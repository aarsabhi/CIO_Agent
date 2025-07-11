import os
import logging
from typing import Dict, List, Optional
import PyPDF2
import docx
import pandas as pd
from openpyxl import load_workbook

logger = logging.getLogger(__name__)

class DocumentProcessor:
    def __init__(self):
        self.supported_extensions = {'.pdf', '.docx', '.xlsx', '.csv', '.txt'}
    
    def process_file(self, file_path: str) -> str:
        """Process uploaded file and extract text content"""
        try:
            file_extension = os.path.splitext(file_path)[1].lower()
            
            if file_extension not in self.supported_extensions:
                raise ValueError(f"Unsupported file type: {file_extension}")
            
            if file_extension == '.pdf':
                return self._process_pdf(file_path)
            elif file_extension == '.docx':
                return self._process_docx(file_path)
            elif file_extension == '.xlsx':
                return self._process_xlsx(file_path)
            elif file_extension == '.csv':
                return self._process_csv(file_path)
            elif file_extension == '.txt':
                return self._process_txt(file_path)
            
        except Exception as e:
            logger.error(f"Error processing file {file_path}: {str(e)}")
            raise
    
    def _process_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                
                return text.strip()
        
        except Exception as e:
            logger.error(f"Error processing PDF {file_path}: {str(e)}")
            return f"Error processing PDF: {str(e)}"
    
    def _process_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = docx.Document(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return text.strip()
        
        except Exception as e:
            logger.error(f"Error processing DOCX {file_path}: {str(e)}")
            return f"Error processing DOCX: {str(e)}"
    
    def _process_xlsx(self, file_path: str) -> str:
        """Extract data from XLSX file"""
        try:
            # First try with pandas (more robust for complex Excel files)
            try:
                logger.info(f"Attempting to process Excel file with pandas: {file_path}")
                df_dict = pd.read_excel(file_path, sheet_name=None, engine='openpyxl')
                
                text = f"Excel File: {os.path.basename(file_path)}\n"
                text += f"Total Sheets: {len(df_dict)}\n\n"
                
                for sheet_name, sheet_df in df_dict.items():
                    text += f"=== Sheet: {sheet_name} ===\n"
                    text += f"Dimensions: {len(sheet_df)} rows × {len(sheet_df.columns)} columns\n"
                    
                    # Clean column names
                    clean_columns = []
                    for col in sheet_df.columns:
                        if pd.isna(col):
                            clean_columns.append("Unnamed_Column")
                        else:
                            clean_columns.append(str(col).strip())
                    
                    text += f"Columns: {', '.join(clean_columns[:10])}"
                    if len(clean_columns) > 10:
                        text += f" ... and {len(clean_columns) - 10} more"
                    text += "\n\n"
                    
                    # Add sample data (first 5 rows)
                    if not sheet_df.empty:
                        text += "Sample Data (first 5 rows):\n"
                        sample_df = sheet_df.head(5).fillna('')
                        
                        # Convert to string representation
                        for idx, row in sample_df.iterrows():
                            row_data = []
                            for col_idx, value in enumerate(row):
                                if col_idx >= 10:  # Limit columns
                                    row_data.append("...")
                                    break
                                row_data.append(str(value)[:50])  # Limit cell content
                            text += f"Row {idx + 1}: {' | '.join(row_data)}\n"
                    else:
                        text += "Sheet is empty\n"
                    
                    text += "\n"
                
                logger.info(f"Successfully processed Excel file with pandas: {file_path}")
                return text.strip()
                
            except Exception as pandas_error:
                logger.warning(f"Pandas failed for {file_path}: {pandas_error}")
                
                # Fallback to openpyxl
                try:
                    logger.info(f"Attempting to process Excel file with openpyxl: {file_path}")
                    workbook = load_workbook(file_path, read_only=True, data_only=True)
                except Exception as openpyxl_error:
                    logger.error(f"Both pandas and openpyxl failed for {file_path}")
                    # Try one more approach - basic file info
                    try:
                        file_size = os.path.getsize(file_path)
                        return f"Excel file detected but could not be processed.\nFile: {os.path.basename(file_path)}\nSize: {file_size} bytes\nError: {str(pandas_error)}\n\nPlease ensure the file is not password-protected, corrupted, or in an unsupported Excel format."
                    except:
                        return f"Excel file processing failed: {str(pandas_error)}"
            
            text = ""
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                text += f"Sheet: {sheet_name}\n"
                
                row_count = 0
                for row in sheet.iter_rows(values_only=True):
                    if row_count > 1000:  # Limit rows to prevent memory issues
                        text += "... (truncated for size)\n"
                        break
                    row_text = " | ".join([str(cell) if cell is not None else "" for cell in row])
                    if row_text.strip():  # Only add non-empty rows
                        text += row_text + "\n"
                    row_count += 1
                
                text += "\n"
            
            workbook.close()
            return text.strip()
        
        except Exception as e:
            logger.error(f"Error processing XLSX {file_path}: {str(e)}")
    
    def _process_csv(self, file_path: str) -> str:
        """Extract data from CSV file"""
        try:
            df = pd.read_csv(file_path)
            
            # Convert DataFrame to text representation
            text = f"CSV Data Summary:\n"
            text += f"Rows: {len(df)}, Columns: {len(df.columns)}\n\n"
            text += f"Columns: {', '.join(df.columns)}\n\n"
            
            # Add first few rows as sample
            text += "Sample Data:\n"
            text += df.head(10).to_string(index=False)
            
            # Add basic statistics for numeric columns
            numeric_cols = df.select_dtypes(include=['number']).columns
            if len(numeric_cols) > 0:
                text += "\n\nNumeric Column Statistics:\n"
                text += df[numeric_cols].describe().to_string()
            
            return text
        
        except Exception as e:
            logger.error(f"Error processing CSV {file_path}: {str(e)}")
            return f"Error processing CSV: {str(e)}"
    
    def _process_txt(self, file_path: str) -> str:
        """Extract text from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    return file.read()
            except Exception as e:
                logger.error(f"Error processing TXT {file_path}: {str(e)}")
                return f"Error processing TXT: {str(e)}"
        
        except Exception as e:
            logger.error(f"Error processing TXT {file_path}: {str(e)}")
            return f"Error processing TXT: {str(e)}"
    
    def get_file_metadata(self, file_path: str) -> Dict:
        """Get metadata about the processed file"""
        try:
            stat = os.stat(file_path)
            return {
                "filename": os.path.basename(file_path),
                "size": stat.st_size,
                "extension": os.path.splitext(file_path)[1].lower(),
                "modified": stat.st_mtime
            }
        
        except Exception as e:
            logger.error(f"Error getting file metadata {file_path}: {str(e)}")
            return {}